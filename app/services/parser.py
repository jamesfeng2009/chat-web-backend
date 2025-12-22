import os
import json
import re

from io import BytesIO
from dataclasses import dataclass
import uuid

from sqlalchemy.orm import Session
import fitz  # PyMuPDF
from docx import Document
import markdown
from bs4 import BeautifulSoup
from paddleocr import PaddleOCR

from app.models.document import Document as DocumentModel
from app.services.document import document_service
from app.core.config import settings
from app.core.logger import get_logger

logger = get_logger(__name__)


@dataclass
class TextBlock:
    """文本块数据结构"""
    text: str
    block_type: str  # paragraph, heading, list, table, etc.
    level: int  # 标题级别或列表级别
    bbox: [Tuple[float, float, float, float]] = None  # (x0, y0, x1, y1)
    page_num: int = 0
    style: [dict[str, any]] = None
    
    def to_dict(self) -> dict[str, any]:
        """转换为字典"""
        return {
            "text": self.text,
            "block_type": self.block_type,
            "level": self.level,
            "bbox": self.bbox,
            "page_num": self.page_num,
            "style": self.style or {}
        }


class BaseParser:
    """基础解析器"""
    
    def parse(self, content: bytes, file_type: str, options: dict[str, any] = None) -> list[TextBlock]:
        """解析文档内容"""
        raise NotImplementedError
    
    def _detect_language(self, text: str) -> str:
        """检测文本语言"""
        # 简单的中文检测
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        if chinese_chars > len(text) * 0.1:  # 如果中文字符占比超过10%
            return "zh"
        return "en"


class PDFParser(BaseParser):
    """PDF解析器"""
    
    def __init__(self):
        self.ocr = None  # 延迟初始化OCR
    
    def _ensure_ocr(self):
        """确保OCR已初始化"""
        if self.ocr is None:
            try:
                self.ocr = PaddleOCR(use_angle_cls=True, lang='ch')
            except Exception as e:
                logger.warning(f"Failed to initialize PaddleOCR: {e}")
                self.ocr = None
    
    def parse(self, content: bytes, file_type: str, options: dict[str, any] = None) -> list[TextBlock]:
        """解析PDF文档"""
        options = options or {}
        use_ocr = options.get("use_ocr", False)
        
        try:
            # 打开PDF文档
            pdf_document = fitz.open(stream=content, filetype="pdf")
            blocks = []
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                
                if use_ocr:
                    # OCR方式解析
                    page_blocks = self._parse_page_with_ocr(page, page_num)
                else:
                    # 文本提取方式解析
                    page_blocks = self._parse_page_with_text(page, page_num)
                
                blocks.extend(page_blocks)
            
            pdf_document.close()
            return blocks
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            raise
    
    def _parse_page_with_text(self, page, page_num: int) -> list[TextBlock]:
        """使用文本提取方式解析页面"""
        blocks = []
        
        # 获取页面文本块
        text_blocks = page.get_text("blocks")
        
        for block in text_blocks:
            # 跳过空白块
            if not block[4].strip():
                continue
            
            # 获取块边界框
            bbox = (block[0], block[1], block[2], block[3])
            
            # 获取文本内容
            text = block[4].strip()
            
            # 判断块类型
            block_type, level = self._detect_block_type(text, page)
            
            # 创建文本块
            text_block = TextBlock(
                text=text,
                block_type=block_type,
                level=level,
                bbox=bbox,
                page_num=page_num,
                style={"font_size": self._estimate_font_size(page, bbox)}
            )
            
            blocks.append(text_block)
        
        # 按位置排序
        blocks.sort(key=lambda b: (b.page_num, b.bbox[1] if b.bbox else 0))
        
        return blocks
    
    def _parse_page_with_ocr(self, page, page_num: int) -> list[TextBlock]:
        """使用OCR方式解析页面"""
        blocks = []
        
        try:
            self._ensure_ocr()
            if not self.ocr:
                logger.warning("OCR not available, falling back to text extraction")
                return self._parse_page_with_text(page, page_num)
            
            # 将页面转换为图片
            pix = page.get_pixmap()
            img_data = pix.tobytes("png")
            
            # OCR识别
            result = self.ocr.ocr(img_data, cls=True)
            
            # 处理OCR结果
            for line in result:
                if line and len(line) > 1:
                    # 获取文本和位置信息
                    text_info = line[1]
                    text = text_info[0]
                    
                    # 获取边界框
                    bbox = line[0]
                    x0 = min(point[0] for point in bbox)
                    y0 = min(point[1] for point in bbox)
                    x1 = max(point[0] for point in bbox)
                    y1 = max(point[1] for point in bbox)
                    
                    # 判断块类型
                    block_type, level = self._detect_block_type(text, page)
                    
                    # 创建文本块
                    text_block = TextBlock(
                        text=text,
                        block_type=block_type,
                        level=level,
                        bbox=(x0, y0, x1, y1),
                        page_num=page_num,
                        style={"confidence": text_info[1] if len(text_info) > 1 else 1.0}
                    )
                    
                    blocks.append(text_block)
            
            # 按位置排序
            blocks.sort(key=lambda b: (b.page_num, b.bbox[1] if b.bbox else 0))
            
        except Exception as e:
            logger.error(f"Error in OCR processing: {e}")
            # 回退到文本提取
            blocks = self._parse_page_with_text(page, page_num)
        
        return blocks
    
    def _detect_block_type(self, text: str, page) -> Tuple[str, int]:
        """检测文本块类型和级别"""
        # 标题检测规则
        heading_patterns = [
            r'^第[一二三四五六七八九十\d]+[章节条款部分]',  # 第一章、第二条等
            r'^[\d]+\.',  # 1.、2.等
            r'^[一二三四五六七八九十]+[、．]',  # 一、二、等
            r'^[（\(]\d+[）\)]',  # (1)、(2)等
            r'^[（\(][一二三四五六七八九十]+[）\)]',  # (一)、(二)等
        ]
        
        for pattern in heading_patterns:
            if re.match(pattern, text):
                # 判断标题级别
                if "第" in pattern and ("章" in pattern or "节" in pattern):
                    return "heading", 1
                elif "条" in pattern:
                    return "heading", 2
                else:
                    return "heading", 3
        
        # 列表检测
        list_patterns = [
            r'^[\d]+\.[\s]',  # 数字列表
            r'^[（\(]\d+[）\)][\s]',  # 括号数字列表
            r'^[•·○●□■▪▫]',  # 符号列表
            r'^[a-zA-Z]\.[\s]',  # 字母列表
        ]
        
        for pattern in list_patterns:
            if re.match(pattern, text):
                return "list", 1
        
        # 表格检测
        if "\t" in text or "|" in text:
            return "table", 1
        
        # 默认为段落
        return "paragraph", 1
    
    def _estimate_font_size(self, page, bbox: Tuple[float, float, float, float]) -> float:
        """估算字体大小"""
        try:
            # 获取文本块内的字体大小
            text = page.get_text("text", clip=bbox)
            if not text:
                return 12.0
            
            # 尝试获取详细的字体信息
            blocks = page.get_text("dict", clip=bbox)
            if blocks and "blocks" in blocks:
                for block in blocks["blocks"]:
                    if "lines" in block:
                        for line in block["lines"]:
                            if "spans" in line:
                                for span in line["spans"]:
                                    if "size" in span:
                                        return span["size"]
            
            # 默认字体大小
            return 12.0
        except:
            return 12.0


class DocxParser(BaseParser):
    """DOCX解析器"""
    
    def parse(self, content: bytes, file_type: str, options: dict[str, any] = None) -> list[TextBlock]:
        """解析DOCX文档"""
        try:
            # 打开DOCX文档
            doc = Document(BytesIO(content))
            blocks = []
            
            # 解析段落
            for i, paragraph in enumerate(doc.paragraphs):
                if not paragraph.text.strip():
                    continue
                
                # 获取段落样式信息
                style_name = paragraph.style.name if paragraph.style else "Normal"
                block_type, level = self._detect_block_type(paragraph.text, style_name)
                
                # 获取字体大小
                font_size = 12.0
                if paragraph.runs:
                    for run in paragraph.runs:
                        if run.font.size:
                            font_size = run.font.size
                            break
                
                # 创建文本块
                text_block = TextBlock(
                    text=paragraph.text.strip(),
                    block_type=block_type,
                    level=level,
                    page_num=0,  # DOCX没有页码概念
                    style={
                        "style_name": style_name,
                        "font_size": font_size,
                        "bold": any(run.bold for run in paragraph.runs if run.bold),
                        "italic": any(run.italic for run in paragraph.runs if run.italic)
                    }
                )
                
                blocks.append(text_block)
            
            # 解析表格
            for table in doc.tables:
                table_text = self._parse_table(table)
                if table_text:
                    text_block = TextBlock(
                        text=table_text,
                        block_type="table",
                        level=1,
                        page_num=0,
                        style={"rows": len(table.rows), "cols": len(table.columns)}
                    )
                    blocks.append(text_block)
            
            return blocks
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise
    
    def _detect_block_type(self, text: str, style_name: str) -> Tuple[str, int]:
        """检测文本块类型和级别"""
        # 根据样式名称判断
        if "Heading" in style_name:
            level = int(style_name.split()[-1]) if style_name.split()[-1].isdigit() else 1
            return "heading", level
        
        # 根据内容判断
        heading_patterns = [
            r'^第[一二三四五六七八九十\d]+[章节条款部分]',
            r'^[\d]+\.',
            r'^[一二三四五六七八九十]+[、．]',
            r'^[（\(]\d+[）\)]',
            r'^[（\(][一二三四五六七八九十]+[）\)]',
        ]
        
        for pattern in heading_patterns:
            if re.match(pattern, text):
                if "第" in pattern and ("章" in pattern or "节" in pattern):
                    return "heading", 1
                elif "条" in pattern:
                    return "heading", 2
                else:
                    return "heading", 3
        
        # 列表检测
        if re.match(r'^[\d]+\.[\s]|[（\(]\d+[）\)][\s]|[•·○●□■▪▫]|[a-zA-Z]\.[\s]', text):
            return "list", 1
        
        # 默认为段落
        return "paragraph", 1
    
    def _parse_table(self, table) -> str:
        """解析表格内容"""
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(cell.text.strip())
            rows.append(" | ".join(cells))
        
        return "\n".join(rows)


class TxtParser(BaseParser):
    """TXT解析器"""
    
    def parse(self, content: bytes, file_type: str, options: dict[str, any] = None) -> list[TextBlock]:
        """解析TXT文档"""
        try:
            # 解码文本
            encoding = options.get("encoding", "utf-8")
            text = content.decode(encoding, errors="replace")
            
            # 分割段落
            paragraphs = text.split("\n\n")
            blocks = []
            
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                
                # 检测块类型
                block_type, level = self._detect_block_type(para)
                
                # 创建文本块
                text_block = TextBlock(
                    text=para,
                    block_type=block_type,
                    level=level,
                    page_num=0
                )
                
                blocks.append(text_block)
            
            return blocks
            
        except Exception as e:
            logger.error(f"Error parsing TXT: {e}")
            raise
    
    def _detect_block_type(self, text: str) -> Tuple[str, int]:
        """检测文本块类型和级别"""
        # 标题检测
        heading_patterns = [
            r'^第[一二三四五六七八九十\d]+[章节条款部分]',
            r'^[\d]+\.',
            r'^[一二三四五六七八九十]+[、．]',
            r'^[（\(]\d+[）\)]',
            r'^[（\(][一二三四五六七八九十]+[）\)]',
        ]
        
        for pattern in heading_patterns:
            if re.match(pattern, text):
                if "第" in pattern and ("章" in pattern or "节" in pattern):
                    return "heading", 1
                elif "条" in pattern:
                    return "heading", 2
                else:
                    return "heading", 3
        
        # 列表检测
        if re.match(r'^[\d]+\.[\s]|[（\(]\d+[）\)][\s]|[•·○●□■▪▫]|[a-zA-Z]\.[\s]', text):
            return "list", 1
        
        # 默认为段落
        return "paragraph", 1


class MarkdownParser(BaseParser):
    """Markdown解析器"""
    
    def parse(self, content: bytes, file_type: str, options: dict[str, any] = None) -> list[TextBlock]:
        """解析Markdown文档"""
        try:
            # 解码文本
            encoding = options.get("encoding", "utf-8")
            text = content.decode(encoding, errors="replace")
            
            # 使用markdown库解析
            md = markdown.Markdown(extensions=['markdown.extensions.tables'])
            html = md.convert(text)
            
            # 使用BeautifulSoup解析HTML
            soup = BeautifulSoup(html, 'html.parser')
            blocks = []
            
            # 解析各种元素
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'table', 'blockquote']):
                tag = element.name
                text = element.get_text().strip()
                
                if not text:
                    continue
                
                # 确定块类型和级别
                if tag.startswith('h'):
                    level = int(tag[1])
                    block_type = "heading"
                elif tag in ['ul', 'ol']:
                    block_type = "list"
                    level = 1
                elif tag == 'table':
                    block_type = "table"
                    level = 1
                elif tag == 'blockquote':
                    block_type = "quote"
                    level = 1
                else:
                    block_type = "paragraph"
                    level = 1
                
                # 创建文本块
                text_block = TextBlock(
                    text=text,
                    block_type=block_type,
                    level=level,
                    page_num=0,
                    style={"html_tag": tag}
                )
                
                blocks.append(text_block)
            
            return blocks
            
        except Exception as e:
            logger.error(f"Error parsing Markdown: {e}")
            raise


class HTMLParser(BaseParser):
    """HTML解析器"""
    
    def parse(self, content: bytes, file_type: str, options: dict[str, any] = None) -> list[TextBlock]:
        """解析HTML文档"""
        try:
            # 解码文本
            encoding = options.get("encoding", "utf-8")
            text = content.decode(encoding, errors="replace")
            
            # 使用BeautifulSoup解析HTML
            soup = BeautifulSoup(text, 'html.parser')
            blocks = []
            
            # 移除脚本和样式
            for script in soup(["script", "style"]):
                script.decompose()
            
            # 解析各种元素
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'table', 'div', 'blockquote']):
                tag = element.name
                text = element.get_text().strip()
                
                if not text:
                    continue
                
                # 确定块类型和级别
                if tag.startswith('h'):
                    level = int(tag[1])
                    block_type = "heading"
                elif tag in ['ul', 'ol']:
                    block_type = "list"
                    level = 1
                elif tag == 'table':
                    block_type = "table"
                    level = 1
                elif tag == 'blockquote':
                    block_type = "quote"
                    level = 1
                else:
                    block_type = "paragraph"
                    level = 1
                
                # 创建文本块
                text_block = TextBlock(
                    text=text,
                    block_type=block_type,
                    level=level,
                    page_num=0,
                    style={"html_tag": tag}
                )
                
                blocks.append(text_block)
            
            return blocks
            
        except Exception as e:
            logger.error(f"Error parsing HTML: {e}")
            raise


class ParserService:
    """文档解析服务"""
    
    def __init__(self):
        self.parsers = {
            "pdf": PDFParser(),
            "docx": DocxParser(),
            "txt": TxtParser(),
            "md": MarkdownParser(),
            "html": HTMLParser()
        }
    
    def parse_document(
        self,
        db: Session,
        document_id: str,
        parser_type: str = "auto",
        options: dict[str, any] = None
    ) -> dict[str, any]:
        """
        解析文档
        
        Args:
            db: 数据库会话
            document_id: 文档ID
            parser_type: 解析器类型
            options: 解析选项
            
        Returns:
            解析结果
        """
        try:
            # 获取文档信息
            document = document_service.get_document(db, document_id=document_id)
            if not document:
                raise ValueError(f"Document not found: {document_id}")
            
            file_type = document.get("file_type")
            
            # 自动选择解析器
            if parser_type == "auto":
                parser_type = file_type
            
            # 获取解析器
            parser = self.parsers.get(parser_type)
            if not parser:
                raise ValueError(f"Unsupported parser type: {parser_type}")
            
            # 获取文件内容
            file_content = document_service.get_file_content(db, document_id=document_id)
            if not file_content:
                raise ValueError(f"File content not found for document: {document_id}")
            
            # 解析文档
            blocks = parser.parse(file_content, file_type, options)
            
            # 保存解析结果
            result = {
                "document_id": document_id,
                "file_type": file_type,
                "parser_type": parser_type,
                "options": options or {},
                "blocks": [block.to_dict() for block in blocks],
                "total_blocks": len(blocks),
                "language": self._detect_document_language(blocks)
            }
            
            # 存储解析结果
            self._save_parse_result(db, document_id, result)
            
            # 更新文档状态
            document_service.update_document_status(
                db=db,
                document_id=document_id,
                parse_status="completed"
            )
            
            return result
            
        except Exception as e:
            logger.error(f"Error parsing document: {e}")
            # 更新文档状态为失败
            document_service.update_document_status(
                db=db,
                document_id=document_id,
                parse_status="failed"
            )
            raise
    
    def get_parse_result(self, db: Session, document_id: str) -> [dict[str, any]]:
        """
        获取解析结果
        
        Args:
            db: 数据库会话
            document_id: 文档ID
            
        Returns:
            解析结果或None
        """
        # 这里可以从数据库或文件系统中获取已保存的解析结果
        # 为了简化，这里返回None，实际实现中可以存储解析结果
        return None
    
    def _detect_document_language(self, blocks: list[TextBlock]) -> str:
        """检测文档语言"""
        if not blocks:
            return "unknown"
        
        # 合并所有文本
        all_text = " ".join([block.text for block in blocks])
        
        # 检测语言
        if len(blocks) > 0 and hasattr(blocks[0], '_detect_language'):
            return blocks[0]._detect_language(all_text)
        else:
            # 默认中文检测
            chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', all_text))
            if chinese_chars > len(all_text) * 0.1:
                return "zh"
            return "en"
    
    def _save_parse_result(self, db: Session, document_id: str, result: dict[str, any]):
        """保存解析结果"""
        # 这里可以将解析结果保存到数据库或文件系统
        # 为了简化，这里只是记录日志
        logger.info(f"Saved parse result for document {document_id}: {len(result.get('blocks', []))} blocks")


# 全局解析服务实例
parser_service = ParserService()