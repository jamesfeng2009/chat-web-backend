"""
文档路由服务
根据传入参数决定解析路径：URL下载或直接解析富文本内容
"""
import os
import hashlib
import requests

from bs4 import BeautifulSoup
from io import BytesIO
import fitz  # PyMuPDF for PDF
from docx import Document as DocxDocument
import pandas as pd

from app.core.logger import logger
from app.services.storage import StorageService


class DocumentRoutingService:
    """文档路由服务，处理多种文档格式的解析入口"""
    
    def __init__(self, storage_service: StorageService):
        self.storage_service = storage_service
    
    def route_and_parse(self, doc_data: dict[str, any]) -> dict[str, any]:
        """
        路由文档解析流程
        
        Args:
            doc_data: 包含文档信息的字典
                - id: 文档ID
                - title: 文档标题
                - type: 文档类型
                - created_at: 创建时间
                - file_url: 文档下载URL
                - rich_content: 富文本内容
                - drafters: 起草人信息
        
        Returns:
            解析后的文档内容，包括文本段落和元数据
        """
        doc_id = doc_data.get("id")
        title = doc_data.get("title", "")
        doc_type = doc_data.get("type", "")
        created_at = doc_data.get("created_at")
        file_url = doc_data.get("file_url", "")
        rich_content = doc_data.get("rich_content", "")
        drafters = doc_data.get("drafters", [])
        
        # 优先使用富文本内容
        if rich_content:
            logger.info(f"Processing rich content for document {doc_id}")
            content = rich_content
            file_type = self._detect_content_type(content)
            raw_text, metadata = self._parse_content_by_type(content, file_type)
        elif file_url:
            logger.info(f"Downloading and processing document from URL: {file_url}")
            try:
                # 下载文件
                downloaded_content, file_type = self._download_file(file_url, doc_id)
                # 解析文件
                raw_text, metadata = self._parse_file_by_type(downloaded_content, file_type)
            except Exception as e:
                logger.error(f"Failed to download or parse file from URL: {str(e)}")
                raise ValueError(f"文档下载或解析失败: {str(e)}")
        else:
            raise ValueError("必须提供file_url或rich_content中的一个")
        
        # 检查文档类型
        if not file_type:
            file_type = self._infer_document_type(raw_text, title, doc_type)
        
        # 处理段落分割
        parsed_segments = self._segment_text(raw_text)
        
        # 计算校验和
        checksum = self._calculate_checksum(raw_text)
        
        # 构建元数据
        document_metadata = {
            "title": title,
            "type": doc_type,
            "created_at": created_at,
            "drafters": drafters,
            "file_type": file_type,
            "checksum": checksum,
            "source": "rich_content" if rich_content else "file_url",
            "original_metadata": metadata
        }
        
        return {
            "parsed_segments": parsed_segments,
            "metadata": document_metadata
        }
    
    def _detect_content_type(self, content: str) -> [str]:
        """检测富文本内容的类型"""
        if content.strip().startswith("<!DOCTYPE") or content.strip().startswith("<html"):
            return "html"
        elif content.strip().startswith("<p") or "<div" in content:
            return "html"
        else:
            return "text"
    
    def _download_file(self, file_url: str, doc_id: str) -> Tuple[bytes, str]:
        """下载文件并返回文件内容和类型"""
        try:
            response = requests.get(file_url, stream=True, timeout=30)
            response.raise_for_status()
            
            content = response.content
            
            # 尝试从URL或Content-Type获取文件类型
            file_type = None
            content_type = response.headers.get("content-type", "")
            
            if "pdf" in content_type:
                file_type = "pdf"
            elif "word" in content_type or "docx" in content_type:
                file_type = "docx"
            elif "excel" in content_type or "xlsx" in content_type:
                file_type = "xlsx"
            elif "html" in content_type:
                file_type = "html"
            elif "text" in content_type:
                file_type = "txt"
            
            # 如果从Content-Type无法确定，尝试从URL确定
            if not file_type:
                if file_url.endswith(".pdf"):
                    file_type = "pdf"
                elif file_url.endswith(".docx") or file_url.endswith(".doc"):
                    file_type = "docx"
                elif file_url.endswith(".xlsx") or file_url.endswith(".xls"):
                    file_type = "xlsx"
                elif file_url.endswith(".html") or file_url.endswith(".htm"):
                    file_type = "html"
                elif file_url.endswith(".txt") or file_url.endswith(".md"):
                    file_type = "txt"
            
            return content, file_type
        
        except Exception as e:
            logger.error(f"Failed to download file from {file_url}: {str(e)}")
            raise
    
    def _parse_content_by_type(self, content: str, file_type: str) -> Tuple[str, Dict]:
        """根据内容类型解析内容"""
        if file_type == "html":
            return self._parse_html(content)
        else:
            # 对于纯文本，直接返回
            return content, {}
    
    def _parse_file_by_type(self, content: bytes, file_type: str) -> Tuple[str, Dict]:
        """根据文件类型解析文件"""
        try:
            if file_type == "pdf":
                return self._parse_pdf(content)
            elif file_type == "docx":
                return self._parse_docx(content)
            elif file_type == "xlsx":
                return self._parse_xlsx(content)
            elif file_type == "html":
                return self._parse_html(content.decode("utf-8", errors="ignore"))
            else:  # txt or others
                return content.decode("utf-8", errors="ignore"), {}
        except Exception as e:
            logger.error(f"Failed to parse {file_type} file: {str(e)}")
            raise
    
    def _parse_pdf(self, content: bytes) -> Tuple[str, Dict]:
        """解析PDF文件"""
        doc = fitz.open(stream=content, filetype="pdf")
        text_parts = []
        metadata = {"page_count": doc.page_count}
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            text_parts.append(page.get_text())
        
        doc.close()
        return "\n".join(text_parts), metadata
    
    def _parse_docx(self, content: bytes) -> Tuple[str, Dict]:
        """解析DOCX文件"""
        doc = DocxDocument(BytesIO(content))
        text_parts = []
        metadata = {"paragraph_count": len(doc.paragraphs)}
        
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)
        
        # 处理表格
        tables_text = []
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                table_text.append("\t".join(row_text))
            tables_text.append("\n".join(table_text))
        
        # 合并段落和表格
        all_text = text_parts + tables_text
        metadata["table_count"] = len(doc.tables)
        
        return "\n".join(all_text), metadata
    
    def _parse_xlsx(self, content: bytes) -> Tuple[str, Dict]:
        """解析XLSX文件"""
        excel_file = BytesIO(content)
        sheets = pd.read_excel(excel_file, sheet_name=None)
        text_parts = []
        metadata = {"sheet_count": len(sheets)}
        
        for sheet_name, df in sheets.items():
            # 将DataFrame转换为字符串
            sheet_text = df.to_string(index=False)
            text_parts.append(f"=== {sheet_name} ===\n{sheet_text}")
        
        return "\n\n".join(text_parts), metadata
    
    def _parse_html(self, html_content: str) -> Tuple[str, Dict]:
        """解析HTML内容"""
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 移除脚本和样式元素
        for script in soup(["script", "style"]):
            script.extract()
        
        # 获取文本
        text = soup.get_text()
        
        # 清理文本，移除多余的空行
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        # 提取元数据
        title_tag = soup.find('title')
        metadata = {
            "title": title_tag.get_text() if title_tag else "",
            "links": [link.get('href') for link in soup.find_all('a', href=True)]
        }
        
        return text, metadata
    
    def _infer_document_type(self, text: str, title: str, doc_type: str) -> str:
        """从文本内容和标题推断文档类型"""
        # 如果已经提供了文档类型，使用它
        if doc_type:
            return doc_type
        
        # 从标题推断
        title_lower = title.lower()
        if "合同" in title_lower or "协议" in title_lower:
            return "合同"
        elif "章程" in title_lower:
            return "章程"
        elif "条款" in title_lower:
            return "条款"
        
        # 从内容推断
        text_lower = text.lower()
        if any(keyword in text_lower for keyword in ["第一条", "第1条", "1. "]):
            return "合同"
        elif any(keyword in text_lower for keyword in ["附件", "附录"]):
            return "附件"
        
        return "文档"
    
    def _segment_text(self, text: str) -> list[dict]:
        """将文本分割成段落片段"""
        # 简单的段落分割
        lines = text.split('\n')
        segments = []
        
        current_id = 0
        for line in lines:
            line = line.strip()
            if line:  # 跳过空行
                segments.append({
                    "id": current_id,
                    "text": line
                })
                current_id += 1
        
        return segments
    
    def _calculate_checksum(self, content: str) -> str:
        """计算内容的SHA256校验和"""
        return hashlib.sha256(content.encode('utf-8')).hexdigest()